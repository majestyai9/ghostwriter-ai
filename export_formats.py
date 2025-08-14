"""
Multiple Export Formats System - EPUB, PDF, DOCX, HTML
"""
import os
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import re

# Import format-specific libraries with fallback
try:
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown2
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

class BookExporter:
    """Export books to multiple formats"""
    
    def __init__(self, project_dir: Path = None):
        """
        Initialize book exporter
        
        Args:
            project_dir: Project directory
        """
        self.project_dir = project_dir or Path(".")
        self.export_dir = self.project_dir / "exports"
        self.export_dir.mkdir(exist_ok=True)
        self.logger = logging.getLogger(__name__)
        
    def export(self,
              book_data: Dict[str, Any],
              format: str,
              options: Dict[str, Any] = None) -> str:
        """
        Export book to specified format
        
        Args:
            book_data: Book data dictionary
            format: Export format (epub, pdf, docx, html)
            options: Format-specific options
            
        Returns:
            Path to exported file
        """
        options = options or {}
        
        if format.lower() == 'epub':
            return self.export_epub(book_data, options)
        elif format.lower() == 'pdf':
            return self.export_pdf(book_data, options)
        elif format.lower() == 'docx':
            return self.export_docx(book_data, options)
        elif format.lower() == 'html':
            return self.export_html(book_data, options)
        else:
            raise ValueError(f"Unsupported format: {format}")
            
    def export_epub(self, 
                   book_data: Dict[str, Any],
                   options: Dict[str, Any]) -> str:
        """
        Export to EPUB format
        
        Args:
            book_data: Book data
            options: EPUB options
            
        Returns:
            Path to EPUB file
        """
        if not EPUB_AVAILABLE:
            raise ImportError("ebooklib not installed. Run: pip install ebooklib")
            
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(f"ghostwriter-{datetime.now().timestamp()}")
        book.set_title(book_data.get('title', 'Untitled'))
        book.set_language(book_data.get('language', 'en'))
        
        # Set optional metadata
        if options.get('author'):
            book.add_author(options['author'])
        if options.get('cover_image'):
            book.set_cover('cover.jpg', open(options['cover_image'], 'rb').read())
            
        # Create chapters
        epub_chapters = []
        spine = ['nav']
        
        # Add introduction/summary if exists
        if book_data.get('summary'):
            intro = epub.EpubHtml(
                title='Introduction',
                file_name='intro.xhtml',
                lang=book_data.get('language', 'en')
            )
            intro.content = f"<h1>Introduction</h1><p>{book_data['summary']}</p>"
            book.add_item(intro)
            spine.append(intro)
            
        # Add chapters
        toc = book_data.get('toc', {})
        for chapter in toc.get('chapters', []):
            ch = epub.EpubHtml(
                title=f"{chapter['number']}. {chapter['title']}",
                file_name=f"chapter_{chapter['number']}.xhtml",
                lang=book_data.get('language', 'en')
            )
            
            # Convert markdown to HTML
            content_html = self._markdown_to_html(chapter.get('content', ''))
            
            # Add sections
            sections_html = ""
            for section in chapter.get('sections', []):
                section_content = self._markdown_to_html(section.get('content', ''))
                sections_html += f"""
                <h2>{chapter['number']}.{section['number']}. {section['title']}</h2>
                {section_content}
                """
                
            ch.content = f"""
            <h1>{chapter['number']}. {chapter['title']}</h1>
            {content_html}
            {sections_html}
            """
            
            book.add_item(ch)
            epub_chapters.append(ch)
            spine.append(ch)
            
        # Add navigation
        book.toc = epub_chapters
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Set spine
        book.spine = spine
        
        # Write file
        filename = self._sanitize_filename(book_data.get('title', 'book')) + '.epub'
        filepath = self.export_dir / filename
        epub.write_epub(str(filepath), book, {})
        
        self.logger.info(f"Exported EPUB to {filepath}")
        return str(filepath)
        
    def export_pdf(self,
                  book_data: Dict[str, Any],
                  options: Dict[str, Any]) -> str:
        """
        Export to PDF format
        
        Args:
            book_data: Book data
            options: PDF options
            
        Returns:
            Path to PDF file
        """
        if not PDF_AVAILABLE:
            raise ImportError("reportlab not installed. Run: pip install reportlab")
            
        filename = self._sanitize_filename(book_data.get('title', 'book')) + '.pdf'
        filepath = self.export_dir / filename
        
        # Create PDF document
        page_size = A4 if options.get('page_size') == 'A4' else letter
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=page_size,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Justify',
            parent=styles['Normal'],
            alignment=TA_JUSTIFY,
            fontSize=11,
            leading=14
        ))
        styles.add(ParagraphStyle(
            name='ChapterTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=30
        ))
        
        # Build content
        story = []
        
        # Title page
        story.append(Paragraph(book_data.get('title', 'Untitled'), styles['Title']))
        story.append(Spacer(1, 0.5*inch))
        
        if options.get('author'):
            story.append(Paragraph(f"By {options['author']}", styles['Heading2']))
            story.append(Spacer(1, 0.5*inch))
            
        if book_data.get('summary'):
            story.append(Paragraph(book_data['summary'], styles['Justify']))
            
        story.append(PageBreak())
        
        # Table of Contents
        story.append(Paragraph("Table of Contents", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        toc = book_data.get('toc', {})
        for chapter in toc.get('chapters', []):
            toc_entry = f"{chapter['number']}. {chapter['title']}"
            story.append(Paragraph(toc_entry, styles['Normal']))
            
        story.append(PageBreak())
        
        # Chapters
        for chapter in toc.get('chapters', []):
            # Chapter title
            title = f"{chapter['number']}. {chapter['title']}"
            story.append(Paragraph(title, styles['ChapterTitle']))
            
            # Chapter content
            if chapter.get('content'):
                # Clean and split into paragraphs
                paragraphs = self._split_into_paragraphs(chapter['content'])
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para, styles['Justify']))
                        story.append(Spacer(1, 0.1*inch))
                        
            # Sections
            for section in chapter.get('sections', []):
                section_title = f"{chapter['number']}.{section['number']}. {section['title']}"
                story.append(Paragraph(section_title, styles['Heading2']))
                story.append(Spacer(1, 0.1*inch))
                
                if section.get('content'):
                    paragraphs = self._split_into_paragraphs(section['content'])
                    for para in paragraphs:
                        if para.strip():
                            story.append(Paragraph(para, styles['Justify']))
                            story.append(Spacer(1, 0.1*inch))
                            
            story.append(PageBreak())
            
        # Build PDF
        doc.build(story)
        
        self.logger.info(f"Exported PDF to {filepath}")
        return str(filepath)
        
    def export_docx(self,
                   book_data: Dict[str, Any],
                   options: Dict[str, Any]) -> str:
        """
        Export to DOCX format
        
        Args:
            book_data: Book data
            options: DOCX options
            
        Returns:
            Path to DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
            
        document = Document()
        
        # Set document properties
        document.core_properties.title = book_data.get('title', 'Untitled')
        if options.get('author'):
            document.core_properties.author = options['author']
            
        # Title page
        title_para = document.add_heading(book_data.get('title', 'Untitled'), 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if options.get('author'):
            author_para = document.add_paragraph(f"By {options['author']}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        document.add_page_break()
        
        # Summary
        if book_data.get('summary'):
            document.add_heading('Summary', 1)
            document.add_paragraph(book_data['summary'])
            document.add_page_break()
            
        # Table of Contents
        document.add_heading('Table of Contents', 1)
        toc = book_data.get('toc', {})
        
        for chapter in toc.get('chapters', []):
            toc_para = document.add_paragraph()
            toc_para.add_run(f"{chapter['number']}. {chapter['title']}")
            
            # Add section entries
            for section in chapter.get('sections', []):
                section_para = document.add_paragraph()
                section_para.add_run(
                    f"    {chapter['number']}.{section['number']}. {section['title']}"
                )
                
        document.add_page_break()
        
        # Chapters
        for chapter in toc.get('chapters', []):
            # Chapter heading
            document.add_heading(f"{chapter['number']}. {chapter['title']}", 1)
            
            # Chapter content
            if chapter.get('content'):
                paragraphs = self._split_into_paragraphs(chapter['content'])
                for para in paragraphs:
                    if para.strip():
                        doc_para = document.add_paragraph(para)
                        doc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
            # Sections
            for section in chapter.get('sections', []):
                document.add_heading(
                    f"{chapter['number']}.{section['number']}. {section['title']}", 
                    2
                )
                
                if section.get('content'):
                    paragraphs = self._split_into_paragraphs(section['content'])
                    for para in paragraphs:
                        if para.strip():
                            doc_para = document.add_paragraph(para)
                            doc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            
            document.add_page_break()
            
        # Save document
        filename = self._sanitize_filename(book_data.get('title', 'book')) + '.docx'
        filepath = self.export_dir / filename
        document.save(str(filepath))
        
        self.logger.info(f"Exported DOCX to {filepath}")
        return str(filepath)
        
    def export_html(self,
                   book_data: Dict[str, Any],
                   options: Dict[str, Any]) -> str:
        """
        Export to HTML format
        
        Args:
            book_data: Book data
            options: HTML options
            
        Returns:
            Path to HTML file
        """
        # Create HTML structure
        html_content = f"""
        <!DOCTYPE html>
        <html lang="{book_data.get('language', 'en')}">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{book_data.get('title', 'Untitled')}</title>
            <style>
                body {{
                    font-family: Georgia, serif;
                    line-height: 1.6;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    background: #f5f5f5;
                }}
                .container {{
                    background: white;
                    padding: 40px;
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                }}
                h1 {{
                    color: #333;
                    border-bottom: 2px solid #333;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #666;
                    margin-top: 30px;
                }}
                .chapter {{
                    page-break-before: always;
                    margin-top: 50px;
                }}
                .toc {{
                    background: #f9f9f9;
                    padding: 20px;
                    margin: 20px 0;
                }}
                .toc ul {{
                    list-style-type: none;
                }}
                .toc a {{
                    text-decoration: none;
                    color: #333;
                }}
                .toc a:hover {{
                    text-decoration: underline;
                }}
                @media print {{
                    .chapter {{
                        page-break-before: always;
                    }}
                }}
            </style>
        </head>
        <body>
            <div class="container">
        """
        
        # Title
        html_content += f"<h1>{book_data.get('title', 'Untitled')}</h1>"
        
        if options.get('author'):
            html_content += f"<p style='text-align: center; font-style: italic;'>By {options['author']}</p>"
            
        # Summary
        if book_data.get('summary'):
            html_content += f"""
            <div class="summary">
                <h2>Summary</h2>
                <p>{book_data['summary']}</p>
            </div>
            """
            
        # Table of Contents
        html_content += """
        <div class="toc">
            <h2>Table of Contents</h2>
            <ul>
        """
        
        toc = book_data.get('toc', {})
        for chapter in toc.get('chapters', []):
            chapter_id = f"chapter-{chapter['number']}"
            html_content += f"""
            <li>
                <a href="#{chapter_id}">{chapter['number']}. {chapter['title']}</a>
            """
            
            if chapter.get('sections'):
                html_content += "<ul>"
                for section in chapter['sections']:
                    section_id = f"section-{chapter['number']}-{section['number']}"
                    html_content += f"""
                    <li>
                        <a href="#{section_id}">
                            {chapter['number']}.{section['number']}. {section['title']}
                        </a>
                    </li>
                    """
                html_content += "</ul>"
                
            html_content += "</li>"
            
        html_content += """
            </ul>
        </div>
        """
        
        # Chapters
        for chapter in toc.get('chapters', []):
            chapter_id = f"chapter-{chapter['number']}"
            html_content += f"""
            <div class="chapter" id="{chapter_id}">
                <h1>{chapter['number']}. {chapter['title']}</h1>
            """
            
            if chapter.get('content'):
                html_content += self._markdown_to_html(chapter['content'])
                
            # Sections
            for section in chapter.get('sections', []):
                section_id = f"section-{chapter['number']}-{section['number']}"
                html_content += f"""
                <div class="section" id="{section_id}">
                    <h2>{chapter['number']}.{section['number']}. {section['title']}</h2>
                """
                
                if section.get('content'):
                    html_content += self._markdown_to_html(section['content'])
                    
                html_content += "</div>"
                
            html_content += "</div>"
            
        # Close HTML
        html_content += """
            </div>
        </body>
        </html>
        """
        
        # Save file
        filename = self._sanitize_filename(book_data.get('title', 'book')) + '.html'
        filepath = self.export_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
            
        self.logger.info(f"Exported HTML to {filepath}")
        return str(filepath)
        
    def _markdown_to_html(self, text: str) -> str:
        """Convert markdown to HTML"""
        if MARKDOWN_AVAILABLE:
            return markdown2.markdown(text)
        else:
            # Basic conversion
            text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
            text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
            text = text.replace('\n\n', '</p><p>')
            return f"<p>{text}</p>"
            
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """Split text into paragraphs"""
        # Remove markdown formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'#{1,6}\s+', '', text)
        
        # Split by double newlines
        paragraphs = text.split('\n\n')
        
        # Clean each paragraph
        cleaned = []
        for para in paragraphs:
            para = para.strip()
            if para:
                # Replace single newlines with spaces
                para = para.replace('\n', ' ')
                cleaned.append(para)
                
        return cleaned
        
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for filesystem"""
        # Remove invalid characters
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        # Replace spaces with underscores
        filename = filename.replace(' ', '_')
        # Limit length
        if len(filename) > 100:
            filename = filename[:100]
        return filename
        
    def export_all_formats(self,
                          book_data: Dict[str, Any],
                          options: Dict[str, Any] = None) -> Dict[str, str]:
        """
        Export to all available formats
        
        Args:
            book_data: Book data
            options: Export options
            
        Returns:
            Dictionary of format: filepath
        """
        results = {}
        options = options or {}
        
        # Try each format
        formats = ['epub', 'pdf', 'docx', 'html']
        
        for fmt in formats:
            try:
                filepath = self.export(book_data, fmt, options)
                results[fmt] = filepath
            except Exception as e:
                self.logger.error(f"Failed to export {fmt}: {e}")
                results[fmt] = None
                
        return results

# Note: Exporter should be obtained from ProjectManager
# to ensure proper project isolation. Use:
# from project_manager import get_project_manager
# pm = get_project_manager()
# exporter = pm.get_exporter()